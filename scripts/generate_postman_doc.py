from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Inventory API Testing (Postman)', level=1)

doc.add_paragraph('Use the following endpoints to test inventory workflows. Take a screenshot of the request and response in Postman and paste it below each section. Replace the placeholder text with your screenshots once available.')

sections = [
    ('1. Create Product (auto-creates inventory)',
     'POST http://localhost:3000/api/v1/products\n'
     'Body (JSON):\n'
     '{\n'
     '  "title": "Sample Product",\n'
     '  "price": 100,\n'
     '  "description": "Test product",\n'
     '  "category": "<CATEGORY_ID>",\n'
     '  "images": ["https://placeimg.com/640/480/any"]\n'
     '}'),
    ('2. Get All Inventory',
     'GET http://localhost:3000/api/v1/inventory'),
    ('3. Get Inventory by ID (with product populated)',
     'GET http://localhost:3000/api/v1/inventory/<INVENTORY_ID>'),
    ('4. Add Stock',
     'POST http://localhost:3000/api/v1/inventory/add_stock\n'
     'Body (JSON):\n'
     '{\n'
     '  "product": "<PRODUCT_ID>",\n'
     '  "quantity": 10\n'
     '}'),
    ('5. Remove Stock',
     'POST http://localhost:3000/api/v1/inventory/remove_stock\n'
     'Body (JSON):\n'
     '{\n'
     '  "product": "<PRODUCT_ID>",\n'
     '  "quantity": 3\n'
     '}'),
    ('6. Reserve Stock',
     'POST http://localhost:3000/api/v1/inventory/reservation\n'
     'Body (JSON):\n'
     '{\n'
     '  "product": "<PRODUCT_ID>",\n'
     '  "quantity": 2\n'
     '}'),
    ('7. Mark Sold',
     'POST http://localhost:3000/api/v1/inventory/sold\n'
     'Body (JSON):\n'
     '{\n'
     '  "product": "<PRODUCT_ID>",\n'
     '  "quantity": 1\n'
     '}')
]

for title, text in sections:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    for line in text.split('\n'):
        p.add_run(line).font.size = Pt(10)
        p.add_run('\n')
    doc.add_paragraph('Screenshot (paste image here):')
    doc.add_paragraph('')

path = 'Inventory_Postman_Test.docx'
doc.save(path)
print('Created', path)
